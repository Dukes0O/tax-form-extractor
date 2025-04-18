import os
import logging
import tempfile
from docx import Document
import subprocess

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_path):
    """
    Extract text from a Word document
    
    Args:
        docx_path (str): Path to the Word document
    
    Returns:
        str: Extracted text content
    """
    try:
        logger.info(f"Extracting text from Word document: {docx_path}")
        
        # Check if file is .docx or .doc
        file_ext = os.path.splitext(docx_path)[1].lower()
        
        if file_ext == '.docx':
            # Use python-docx for .docx files
            doc = Document(docx_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        
        elif file_ext == '.doc':
            # For older .doc format, we can try to convert to .docx first
            # This requires LibreOffice or similar to be installed
            logger.warning("Older .doc format detected. Attempting conversion.")
            
            temp_docx = os.path.join(tempfile.gettempdir(), f"converted_{os.path.basename(docx_path)}x")
            
            try:
                # Try using LibreOffice to convert (if installed)
                subprocess.run([
                    'soffice', '--headless', '--convert-to', 'docx', 
                    '--outdir', os.path.dirname(temp_docx), docx_path
                ], check=True)
                
                # Now process the converted file
                return extract_text_from_docx(temp_docx)
                
            except (subprocess.SubprocessError, FileNotFoundError):
                logger.error("LibreOffice conversion failed or not installed")
                raise RuntimeError("Cannot process .doc files. LibreOffice not installed or conversion failed.")
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        raise

def convert_docx_to_pdf(docx_path):
    """
    Convert Word document to PDF for further processing
    
    Args:
        docx_path (str): Path to the Word document
    
    Returns:
        str: Path to the generated PDF file
    """
    try:
        logger.info(f"Converting Word document to PDF: {docx_path}")
        
        # Generate output PDF path
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        
        # Use LibreOffice for conversion (if installed)
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', os.path.dirname(pdf_path), docx_path
        ], check=True)
        
        if os.path.exists(pdf_path):
            logger.info(f"Word document converted to PDF: {pdf_path}")
            return pdf_path
        else:
            raise FileNotFoundError(f"PDF conversion failed: {pdf_path} not found")
    
    except Exception as e:
        logger.error(f"Error converting Word document to PDF: {str(e)}")
        raise
